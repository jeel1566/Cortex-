import os
import csv
from typing import List, Dict, Any, Optional
from app.ingestion.connectors.base import ConnectorAdapter
from app.ingestion.engine_models import (
    NormalizedSourceBundle,
    NormalizedSourceObject,
    NormalizedSourceDocument,
    NormalizedSourceSegment,
)

class LocalUploadAdapter(ConnectorAdapter):
    def __init__(self, tenant_id: str, file_path: str, filename: Optional[str] = None):
        self.tenant_id = tenant_id
        self.file_path = file_path
        self.filename = filename or os.path.basename(file_path)

    def discover(self) -> List[Dict[str, Any]]:
        # Single file upload discovery
        ext = os.path.splitext(self.filename)[1].lower()
        return [{
            "external_id": f"upload://{self.filename}",
            "filename": self.filename,
            "file_path": self.file_path,
            "ext": ext,
        }]

    def fetch(self, external_id: str) -> Dict[str, Any]:
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"File not found: {self.file_path}")
        return {
            "external_id": external_id,
            "filename": self.filename,
            "file_path": self.file_path,
        }

    def extract(self, raw_data: Dict[str, Any]) -> Dict[str, Any]:
        file_path = raw_data["file_path"]
        ext = os.path.splitext(file_path)[1].lower()
        
        segments = []
        body_text = ""
        
        if ext in [".md", ".markdown"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                body_text = f.read()
            segments = self._parse_markdown(body_text)
            
        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                body_text = f.read()
            segments = self._parse_txt(body_text)
            
        elif ext == ".html":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                body_text = f.read()
            segments = self._parse_html(body_text)
            
        elif ext == ".docx":
            segments = self._parse_docx(file_path)
            body_text = "\n\n".join(s["text"] for s in segments)
            
        elif ext == ".pdf":
            segments = self._parse_pdf(file_path)
            body_text = "\n\n".join(s["text"] for s in segments)
            
        elif ext == ".csv":
            segments = self._parse_csv(file_path)
            body_text = "\n".join(s["text"] for s in segments)
            
        elif ext == ".xlsx":
            segments = self._parse_xlsx(file_path)
            body_text = "\n".join(s["text"] for s in segments)
            
        else:
            raise ValueError(f"Unsupported file format: {ext}")
            
        return {
            "title": os.path.splitext(self.filename)[0],
            "body_text": body_text,
            "segments": segments,
        }

    def normalize(self) -> NormalizedSourceBundle:
        discovered = self.discover()[0]
        raw_data = self.fetch(discovered["external_id"])
        extracted = self.extract(raw_data)
        
        source_obj = NormalizedSourceObject(
            tenant_id=self.tenant_id,
            connector_type="local_upload",
            external_id=discovered["external_id"],
            object_type="file",
            title=extracted["title"],
            metadata={"filename": self.filename},
        )
        
        doc = NormalizedSourceDocument(
            source_object_external_id=discovered["external_id"],
            title=extracted["title"],
            body_text=extracted["body_text"],
            metadata={"filename": self.filename},
        )
        
        segments = []
        for s in extracted["segments"]:
            segments.append(NormalizedSourceSegment(
                document_ref=discovered["external_id"],
                segment_type=s.get("segment_type", "paragraph"),
                heading_path=s.get("heading_path", []),
                position=s["position"],
                text=s["text"],
                metadata=s.get("metadata", {}),
            ))
            
        return NormalizedSourceBundle(
            tenant_id=self.tenant_id,
            connector_type="local_upload",
            objects=[source_obj],
            documents=[doc],
            segments=segments,
            relationships=[]
        )

    def _parse_markdown(self, text: str) -> List[Dict[str, Any]]:
        lines = text.splitlines()
        segments = []
        current_heading_path = []
        position = 0
        current_block = []
        
        def emit_block():
            nonlocal position
            if not current_block:
                return
            block_text = "\n".join(current_block).strip()
            if block_text:
                segments.append({
                    "segment_type": "paragraph",
                    "heading_path": list(current_heading_path),
                    "position": position,
                    "text": block_text
                })
                position += 1
            current_block.clear()

        for line in lines:
            stripped = line.strip()
            if not stripped:
                emit_block()
                continue
            if stripped.startswith("#"):
                emit_block()
                level = 0
                for char in stripped:
                    if char == "#":
                        level += 1
                    else:
                        break
                heading_title = stripped[level:].strip()
                if len(current_heading_path) >= level:
                    current_heading_path = current_heading_path[:level - 1]
                while len(current_heading_path) < level - 1:
                    current_heading_path.append("")
                current_heading_path.append(heading_title)
                
                segments.append({
                    "segment_type": "heading",
                    "heading_path": list(current_heading_path),
                    "position": position,
                    "text": heading_title,
                    "metadata": {"level": level}
                })
                position += 1
            else:
                current_block.append(line)
        emit_block()
        return segments

    def _parse_txt(self, text: str) -> List[Dict[str, Any]]:
        paragraphs = text.split("\n\n")
        segments = []
        position = 0
        for p in paragraphs:
            p_text = p.strip()
            if p_text:
                segments.append({
                    "segment_type": "paragraph",
                    "heading_path": [],
                    "position": position,
                    "text": p_text
                })
                position += 1
        return segments

    def _parse_html(self, text: str) -> List[Dict[str, Any]]:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(text, "html.parser")
        segments = []
        position = 0
        current_heading_path = []
        
        body = soup.body or soup
        for element in body.find_all(recursive=True):
            if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                level = int(element.name[1])
                heading_title = element.get_text().strip()
                if not heading_title:
                    continue
                if len(current_heading_path) >= level:
                    current_heading_path = current_heading_path[:level - 1]
                while len(current_heading_path) < level - 1:
                    current_heading_path.append("")
                current_heading_path.append(heading_title)
                
                segments.append({
                    "segment_type": "heading",
                    "heading_path": list(current_heading_path),
                    "position": position,
                    "text": heading_title,
                    "metadata": {"level": level}
                })
                position += 1
            elif element.name in ["p", "li", "pre"]:
                text_content = element.get_text().strip()
                if text_content:
                    if not element.find(["p", "li", "pre", "h1", "h2", "h3", "h4", "h5", "h6"]):
                        segments.append({
                            "segment_type": "paragraph",
                            "heading_path": list(current_heading_path),
                            "position": position,
                            "text": text_content
                        })
                        position += 1
        return segments

    def _parse_docx(self, file_path: str) -> List[Dict[str, Any]]:
        import docx
        doc = docx.Document(file_path)
        segments = []
        position = 0
        current_heading_path = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name.lower()
            if style_name.startswith("heading"):
                try:
                    level = int(style_name.replace("heading", "").strip())
                except ValueError:
                    level = 1
                if len(current_heading_path) >= level:
                    current_heading_path = current_heading_path[:level - 1]
                while len(current_heading_path) < level - 1:
                    current_heading_path.append("")
                current_heading_path.append(text)
                
                segments.append({
                    "segment_type": "heading",
                    "heading_path": list(current_heading_path),
                    "position": position,
                    "text": text,
                    "metadata": {"level": level}
                })
                position += 1
            else:
                segments.append({
                    "segment_type": "paragraph",
                    "heading_path": list(current_heading_path),
                    "position": position,
                    "text": text
                })
                position += 1
        return segments

    def _parse_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        segments = []
        position = 0
        has_text = False
        
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                has_text = True
                paragraphs = text.split("\n\n")
                for p in paragraphs:
                    p_text = p.strip()
                    if p_text:
                        segments.append({
                            "segment_type": "paragraph",
                            "heading_path": [],
                            "position": position,
                            "text": p_text,
                            "metadata": {"page": page_num + 1}
                        })
                        position += 1
        if not has_text:
            raise ValueError("PDF contains no extractable text. Scanned or OCR-only PDFs are unsupported.")
        return segments

    def _parse_csv(self, file_path: str) -> List[Dict[str, Any]]:
        segments = []
        position = 0
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            rows = list(reader)
        for row in rows:
            if any(cell.strip() for cell in row):
                text = ", ".join(cell.strip() for cell in row)
                segments.append({
                    "segment_type": "table_row",
                    "heading_path": [],
                    "position": position,
                    "text": text,
                    "metadata": {"row_index": position}
                })
                position += 1
        return segments

    def _parse_xlsx(self, file_path: str) -> List[Dict[str, Any]]:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True)
        segments = []
        position = 0
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            for row_idx, row in enumerate(sheet.iter_rows(values_only=True)):
                row_vals = [str(val).strip() for val in row if val is not None]
                if any(row_vals):
                    text = ", ".join(row_vals)
                    segments.append({
                        "segment_type": "table_row",
                        "heading_path": [sheet_name],
                        "position": position,
                        "text": text,
                        "metadata": {"sheet": sheet_name, "row_index": row_idx}
                    })
                    position += 1
        return segments
