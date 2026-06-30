from __future__ import annotations

import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\Cortex")
OUT = ROOT / "docs" / "10-cortex-knowledge-compiler-implementation-plan.docx"
ARCH_IMAGE = Path(
    r"C:\Users\Ayush Prajapati\.codex\generated_images"
    r"\019f0d66-1412-7663-b0bd-b7a76f23dcf1"
    r"\ig_0056bb7c506645bd016a40e534dfc081918a1dadf242b6ff52.png"
)


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 45)
MUTED = RGBColor(88, 96, 105)
LIGHT_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
GREEN_FILL = "E7F4EC"
RED_FILL = "FCE8E6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                set_cell_margins(row.cells[idx])
                row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header_fill=LIGHT_FILL) -> None:
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
                    run.font.color.rgb = INK
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = DARK_BLUE


def set_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Cortex Knowledge Compiler Plan")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Cortex Knowledge Compiler")
    run.font.name = "Calibri"
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Implementation Plan for Multi-Connector Trusted Knowledge")
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=4, cols=2)
    set_table_width(meta, [1.6, 4.9])
    rows = [
        ("Audience", "Founder, engineer, AI coding assistant, and technical reviewers"),
        ("Purpose", "Turn Cortex from sentence-cluster ingestion into a fast, cited, permission-aware knowledge compiler."),
        ("Initial scope", "Local Upload, Notion, Google Docs, and Slack."),
        ("Created", datetime.date.today().isoformat()),
    ]
    for row, (k, v) in zip(meta.rows, rows):
        row.cells[0].text = k
        row.cells[1].text = v
    style_table(meta)

    add_callout(
        doc,
        "Core decision",
        "Keep classifier and clusterer, but remove them from page authorship. They become enrichment, relationship, duplicate, and conflict engines. Cortex's main authoring unit becomes Document -> Section -> Segment -> Proposition -> Page.",
        BLUE_FILL,
    )


def add_callout(doc: Document, label: str, text: str, fill: str = CALLOUT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.rows[0].cells[0].width = Inches(6.5)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = INK
    doc.add_paragraph()


def add_bullets(doc: Document, items) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_code(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.rows[0].cells[0].width = Inches(6.5)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F7F7F7")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(40, 40, 40)
    doc.add_paragraph()


def add_contents(doc: Document) -> None:
    doc.add_heading("Contents", level=1)
    add_numbered(
        doc,
        [
            "Executive summary",
            "Previous architecture review context",
            "Q&A summary of the architecture chat",
            "What changes and why",
            "How the new pipeline works",
            "Target architecture and data model",
            "Connector plan and roadmap",
            "Downsides and solutions",
            "Metrics and acceptance criteria",
            "Future CLI, MCP, sandbox, and dynamic tools",
            "AI coding assistant prompts",
            "Appendix: implementation notes",
        ],
    )


def add_executive_summary(doc: Document) -> None:
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "Cortex should evolve from a RAG-like ingestion pipeline into a trust-first knowledge compiler. The current center of gravity is sentence splitting, classification, clustering, synthesis, validation, and immediate Git commit. That has helped prove the shape of the product, but it is not the right long-term engine for a system that should behave like a company memory book."
    )
    doc.add_paragraph(
        "The improved version keeps raw source truth, compiles structured pages in the background, requires evidence for claims, and answers with citations under department and clearance controls. It gives users fast search immediately while deeper canonical pages become available after validation and approval."
    )
    add_callout(
        doc,
        "One-sentence goal",
        "Cortex turns messy company sources into structured, cited, connected, permission-aware knowledge that agents and humans can trust.",
        GREEN_FILL,
    )
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Area"
    hdr[1].text = "Current"
    hdr[2].text = "Improved"
    rows = [
        ("Main unit", "Sentence cluster", "Document -> Section -> Segment -> Proposition"),
        ("Speed", "Waits on full pipeline", "Raw segments searchable first"),
        ("Trust", "Generated page can become truth too early", "Only approved pages enter Git"),
        ("Connectors", "Notion/Slack-shaped messages", "Adapter output normalized source objects"),
        ("PII", "Redaction disabled", "Data preserved, retrieval access-filtered"),
        ("RAG comparison", "Likely chunks or pages", "Cited pages plus raw evidence plus graph context"),
    ]
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    set_table_width(table, [1.35, 2.45, 2.7])
    style_table(table)


def add_previous_review_context(doc: Document) -> None:
    doc.add_heading("2. Previous Architecture Review Context", level=1)
    doc.add_paragraph(
        "This plan is based on the earlier read-only review thread titled Review Cortex architecture. That review looked at Cortex as a product and architecture, not just as a Notion bug. The conclusion was blunt: Cortex has a strong thesis, but the current Phase 2 implementation is scaffold-complete rather than trust-complete."
    )
    add_callout(
        doc,
        "Original verdict",
        "Cortex can win only if it becomes a verified evidence compiler, not just a summarizer or a prettier RAG wrapper.",
        GREEN_FILL,
    )

    doc.add_heading("Critical Findings From The Review", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Finding"
    table.rows[0].cells[1].text = "Why it matters"
    table.rows[0].cells[2].text = "Plan response"
    rows = [
        (
            "Auth fallback is unsafe",
            "Unsigned token fallback and default admin behavior can turn configuration mistakes into full access.",
            "Keep auth hardening as a prerequisite before production use: no unsigned tokens outside explicit dev mode, no default admin, validate issuer/expiry/audience.",
        ),
        (
            "Tenant IDs are trusted in filesystem paths",
            "Tenant isolation depends on path safety. Raw token claims should not be interpolated into paths without canonical validation.",
            "Add tenant/page ID validation and path containment checks before any file, Git, upload, or tenant DB access.",
        ),
        (
            "Claim redaction is fragile",
            "Redacting proposition text while returning the full body can leak paraphrased or repeated sensitive content.",
            "Move to evidence-level permission filtering before answer generation; do not send unauthorized evidence to the LLM.",
        ),
        (
            "Query API discards the generated answer",
            "The query engine can generate an answer, but the API returns placeholder page content instead. This breaks the product promise.",
            "Make query response the product: answer, citations, pages read, source segments read, redactions, confidence, and knowledge gaps.",
        ),
        (
            "Approval inbox is presentation-only",
            "The UI has local alerts but no durable backend draft/approval state.",
            "Add KnowledgePageDraft, approval endpoints, approval records, and Git commit only after approval.",
        ),
        (
            "Rate limiting is not actually implemented",
            "The product claims rate limits, but routes do not enforce tenant or agent quotas.",
            "Add per-tenant and per-agent rate limits before exposing broad connector sync or query APIs.",
        ),
        (
            "Git policy conflicts with trust architecture",
            "Drafts and fallback pages can be committed, polluting the truth layer.",
            "Use Git only for approved canonical pages. Store drafts, rejected outputs, raw connector data, and embeddings outside Git.",
        ),
        (
            "Current Cortex lost to baseline RAG in benchmark",
            "The earlier review noted Cortex v1 averaged 3.00/5 versus RAG at 3.38/5, with RAG winning 17 cases and Cortex winning 3.",
            "Retain raw evidence, add proposition/source links, and make evals a release gate.",
        ),
    ]
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    set_table_width(table, [1.6, 2.45, 2.45])
    style_table(table, header_fill=BLUE_FILL)

    doc.add_heading("Root Cause From The Flowgent/Notion Investigation", level=2)
    doc.add_paragraph(
        "The earlier thread found that the Flowgent document did not appear in saved tenant data. Generated pages were sourced from system sync fallback text rather than notion:// or document:// Flowgent sources. Some pages also included prompt leakage such as Input JSON, Expected Output, and </output_format>. That means the problem was not only page quality; Cortex was often compiling the wrong input and then committing malformed output."
    )
    add_bullets(
        doc,
        [
            "Wrong input: the intended Flowgent document was not ingested.",
            "Silent fallback: empty connector results became demo/system pages.",
            "Bad chunking: sentence clusters created shallow pages.",
            "Weak output gate: prompt-leaked LLM output could be saved.",
            "Weak source model: there was no durable raw source segment store to ground pages.",
        ],
    )

    doc.add_heading("Original Winning Architecture Against RAG", level=2)
    add_code(
        doc,
        "Raw source event\n"
        "  -> raw immutable segment store\n"
        "  -> proposition extraction\n"
        "  -> proposition -> source_segment evidence links\n"
        "  -> canonical page revision\n"
        "  -> validation and approval gate\n"
        "  -> page index + proposition index + raw segment index\n"
        "  -> hybrid query: graph + page vectors + raw evidence vectors\n"
        "  -> answer with citations, confidence, redactions, and feedback hook",
    )

    doc.add_heading("Original Improvement Sequence", level=2)
    add_numbered(
        doc,
        [
            "Security hardening: signed auth, validated tenant IDs, path containment, no default admin.",
            "Truth model: raw segments, propositions, evidence links, page state, approval records, source access levels.",
            "Retrieval rewrite: hybrid page plus raw evidence retrieval with graph traversal and vector safety net.",
            "Approval workflow: durable draft/pending/approved/rejected/conflicted states and real inbox endpoints.",
            "Evaluation gate: benchmark Cortex vs RAG on accuracy, completeness, citations, leakage, and latency.",
            "Connector reality: replace mock connector paths with adapter contracts, backoff, idempotency, and checkpointing.",
            "CLI and MCP: ship read-first interfaces after the trust layer is solid.",
            "Dynamic tools and sandbox: ship last, behind capability policies, sandbox execution, and human approval.",
        ],
    )


def add_qa_context(doc: Document) -> None:
    doc.add_heading("3. Q&A Summary For Context", level=1)
    qas = [
        (
            "What is the core architectural change?",
            "Cortex should stop making pages directly from sentence clusters. The new core is a compiler: source documents become sections, sections become source segments, source segments support propositions, and validated propositions become canonical knowledge pages.",
        ),
        (
            "Do classifier and clusterer get deleted?",
            "No. They are demoted. Classifier labels sections, propositions, and pages. Clusterer finds related documents, duplicate topics, conflicts, and missing links. They become librarian tools, not the author of the book.",
        ),
        (
            "Why is Cortex like a book?",
            "A useful knowledge system preserves chapter structure, headings, context, examples, decisions, and relationships. A book is not written by randomly grouping similar sentences. Cortex should compile structured memory, not just retrieve fragments.",
        ),
        (
            "Why is this better than basic RAG?",
            "Basic RAG often retrieves likely chunks. Cortex should return a cited answer, the canonical approved page, exact source segments, propositions, graph neighbors, redactions, confidence, and knowledge gaps.",
        ),
        (
            "Is this only for Notion?",
            "No. Notion is the first visible case, but the architecture is for Local Upload, Notion, Google Docs, Slack, and later Jira, Sheets, GitHub, Discord, and other connectors.",
        ),
        (
            "Will every connector need a different pipeline?",
            "Each connector needs a different adapter, but the compiler should be shared. Slack extracts threads and messages. Google Docs extracts headings and paragraphs. Notion extracts pages and blocks. After normalization, they all produce SourceObject, SourceDocument, SourceSegment, and SourceRelationship records.",
        ),
        (
            "Why not remove PII during ingestion?",
            "Some sensitive data is the knowledge users need, such as lead phone numbers, salary data, pricing, or contract notes. Cortex should preserve it and filter at retrieval time using department, role, and clearance.",
        ),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Question"
    table.rows[0].cells[1].text = "Answer"
    for q, a in qas:
        cells = table.add_row().cells
        cells[0].text = q
        cells[1].text = a
    set_table_width(table, [2.15, 4.35])
    style_table(table, header_fill=BLUE_FILL)


def add_changes(doc: Document) -> None:
    doc.add_heading("4. What Will Change", level=1)
    add_bullets(
        doc,
        [
            "Replace the sentence-cluster main path with a normalized source layer and structure-aware compiler.",
            "Add first-class source storage before any LLM synthesis.",
            "Make Local Upload, Notion, Google Docs, and Slack produce the same normalized output contract.",
            "Add a fast raw-segment search path so users can search data before page compilation finishes.",
            "Add a background compiler that produces evidence-linked propositions and page drafts.",
            "Keep Git only for approved canonical pages; store raw data, drafts, rejected outputs, embeddings, and logs outside Git.",
            "Use permission-aware retrieval instead of destructive PII removal.",
        ],
    )
    doc.add_heading("Current Pipeline", level=2)
    add_code(doc, "raw messages -> sentence split -> classifier -> clusterer -> synthesize page -> validate -> write Markdown -> commit to Git")
    doc.add_heading("Target Pipeline", level=2)
    add_code(
        doc,
        "connector adapter -> normalized source store -> raw segment index\n"
        "                  -> background compiler -> propositions -> page drafts\n"
        "                  -> strict validator -> approval -> Git truth store\n"
        "                  -> hybrid retrieval -> cited answer",
    )


def add_how_it_works(doc: Document) -> None:
    doc.add_heading("5. How The Improved Version Works", level=1)
    doc.add_heading("4.1 Fast Path", level=2)
    doc.add_paragraph(
        "The fast path exists so Cortex feels responsive. As soon as a file is uploaded or a connector sync completes fetching, Cortex stores raw source documents and segments, computes embeddings for segments, and makes them searchable. Page synthesis can lag behind without blocking search."
    )
    add_bullets(
        doc,
        [
            "Target: first searchable results in under 30 seconds for ordinary uploads and small connector syncs.",
            "Skip unchanged records with content hashes.",
            "Use batch embeddings and store raw source segment IDs for citations.",
        ],
    )
    doc.add_heading("4.2 Background Compiler", level=2)
    doc.add_paragraph(
        "The background compiler converts raw source into higher-quality knowledge. It groups segments by document structure or thread, extracts evidence-linked propositions, synthesizes page drafts, validates shape and citations, and pushes drafts into approval."
    )
    add_bullets(
        doc,
        [
            "Document structure strategy: Local Upload, Notion, Google Docs.",
            "Conversation thread strategy: Slack.",
            "Tabular strategy: later for Sheets and CSV-heavy workflows.",
            "Code artifact strategy: later for GitHub repositories and PRs.",
        ],
    )
    doc.add_heading("4.3 Trusted Query", level=2)
    doc.add_paragraph(
        "Query combines approved pages, propositions, raw segments, vector hits, and graph neighbors. Before an answer is generated, candidate evidence is filtered by the user's department, role, and clearance."
    )
    add_code(
        doc,
        "query -> approved page search -> proposition search -> raw segment search -> graph traversal\n"
        "      -> permission filter -> redaction -> answer with citations, confidence, and knowledge gaps",
    )


def add_architecture(doc: Document) -> None:
    doc.add_heading("6. Target Architecture", level=1)
    if ARCH_IMAGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(ARCH_IMAGE), width=Inches(6.5))
    else:
        add_callout(doc, "Diagram missing", "The generated architecture image was not found. Recreate from the text flow below.", RED_FILL)
    doc.add_paragraph(
        "The system is deliberately split into connector adapters, normalized source storage, fast search, background compilation, Git truth, and permission-aware retrieval. This keeps connectors flexible while preserving one shared compiler."
    )
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Layer"
    table.rows[0].cells[1].text = "Responsibility"
    table.rows[0].cells[2].text = "Notes"
    rows = [
        ("Connector adapters", "Discover, fetch, extract, normalize", "One adapter per connector; same output contract."),
        ("Normalized source store", "Store raw truth and metadata", "Never rely only on generated Markdown."),
        ("Raw segment index", "Fast search and fallback evidence", "Available before compilation completes."),
        ("Compiler", "Segment, extract propositions, synthesize drafts", "LLM used after programmatic grouping."),
        ("Validator", "Reject malformed or uncited output", "No prompt leakage; no evidence-free claims."),
        ("Approval and Git", "Promote trusted pages", "Git stores approved canonical Markdown only."),
        ("Permission-aware query", "Filter evidence before answer generation", "Preserve sensitive data but restrict retrieval."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    set_table_width(table, [1.55, 2.45, 2.5])
    style_table(table, header_fill=BLUE_FILL)


def add_data_model(doc: Document) -> None:
    doc.add_heading("7. Normalized Data Model", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Entity"
    table.rows[0].cells[1].text = "Like I am 5"
    table.rows[0].cells[2].text = "Implementation meaning"
    rows = [
        ("SourceObject", "The outside thing Cortex found.", "A Notion page, Slack thread, Google Doc, or uploaded file with external ID, URL, timestamps, raw JSON, and connector type."),
        ("SourceDocument", "The readable version of that thing.", "Extracted title, body text, source object link, content hash, and parsing metadata."),
        ("SourceSegment", "A small piece of the readable thing.", "A heading section, paragraph, table row, Slack message, or code block used as evidence."),
        ("SourceRelationship", "A string connecting two things.", "Links such as Slack thread discusses Google Doc, Notion page links child page, or upload duplicates source document."),
        ("Proposition", "One fact Cortex believes.", "A single claim with confidence, sensitivity, and required evidence segment IDs."),
        ("KnowledgePageDraft", "A page before it is trusted.", "Generated Markdown plus YAML, validation results, source references, and status."),
        ("SyncRun", "A receipt for a sync.", "Counts, failures, connector type, duration, skipped duplicates, and stage status."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    set_table_width(table, [1.45, 2.05, 3.0])
    style_table(table)
    doc.add_heading("Minimum Shared Fields", level=2)
    add_code(
        doc,
        "source_objects: id, tenant_id, connector_type, external_id, object_type, title, url, author, created_at, updated_at, raw_json, content_hash\n"
        "source_documents: id, source_object_id, title, body_text, metadata_json, content_hash\n"
        "source_segments: id, document_id, segment_type, heading_path, position, text, author, timestamp, metadata_json, content_hash\n"
        "source_relationships: id, from_object_id, to_object_id, relationship_type, metadata_json\n"
        "propositions: id, page_id, text, confidence, sensitivity, evidence_segment_ids, validation_status",
    )


def add_connector_plan(doc: Document) -> None:
    doc.add_heading("8. Connector Implementation Plan", level=1)
    table = doc.add_table(rows=1, cols=4)
    for idx, text in enumerate(["Connector", "Adapter extracts", "Segmentation strategy", "V1 parser/API"]):
        table.rows[0].cells[idx].text = text
    rows = [
        ("Local Upload", "Files and parsed text", "Document structure", "Markdown/TXT built-in, pypdf/pdfplumber, python-docx, pandas/openpyxl, BeautifulSoup."),
        ("Notion", "Pages, databases, blocks, child pages", "Document structure", "Notion API search plus recursive block children."),
        ("Google Docs", "Docs, headings, paragraphs, tables", "Document structure", "Google Drive/Docs API export or document structure fetch."),
        ("Slack", "Channels, threads, messages, files", "Conversation thread", "Slack Web API conversations history/replies."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    set_table_width(table, [1.2, 1.95, 1.45, 1.9])
    style_table(table, header_fill=BLUE_FILL)
    add_callout(
        doc,
        "Normalization rule",
        "The adapter is connector-specific. The compiler is shared. Do not build a different end-to-end pipeline for every connector.",
        GREEN_FILL,
    )


def add_implementation_roadmap(doc: Document) -> None:
    doc.add_heading("9. Implementation Roadmap", level=1)
    phases = [
        ("0", "Document and alignment", "Use this plan as the handoff contract. Freeze v1 connector scope."),
        ("1", "Normalized source schema", "Add source objects, documents, segments, relationships, propositions, drafts, and sync runs."),
        ("2", "Local Upload v1", "Parse md/txt/pdf/docx/csv/xlsx/html and index raw segments fast."),
        ("3", "Notion migration", "Change Notion output from message dicts to normalized source records."),
        ("4", "Google Docs adapter", "Extract document tree into source documents and heading-aware segments."),
        ("5", "Slack adapter", "Group channels, threads, messages, and replies into conversation documents."),
        ("6", "Fast path plus compiler", "Raw segment index first; background propositions and drafts second."),
        ("7", "Strict validation", "Reject invalid YAML, prompt leakage, uncited propositions, and evidence-free claims."),
        ("8", "Approval and Git truth", "Only approved pages are written and committed to Git."),
        ("9", "Permission-aware retrieval", "Filter by department, role, clearance, sensitivity, and source permissions."),
        ("10", "Eval against RAG", "Measure answer quality, citations, detail retention, hallucinations, and speed."),
    ]
    table = doc.add_table(rows=1, cols=3)
    for idx, text in enumerate(["Phase", "Name", "Outcome"]):
        table.rows[0].cells[idx].text = text
    for values in phases:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    set_table_width(table, [0.7, 1.9, 3.9])
    style_table(table)


def add_downsides(doc: Document) -> None:
    doc.add_heading("10. Downsides And Solutions", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Downside"
    table.rows[0].cells[1].text = "Risk"
    table.rows[0].cells[2].text = "Solution"
    rows = [
        ("Complexity", "Too much system before product proof.", "Keep modular monolith, narrow connectors to Local Upload, Notion, Google Docs, Slack, and ship in phases."),
        ("Slower than RAG", "Full compilation takes longer than chunk/embed/search.", "Make raw segment search available first; compile trusted pages in background."),
        ("Normalization difficulty", "Each app has a different shape.", "Use a small shared schema plus connector-specific metadata_json."),
        ("Expensive propositions", "LLM calls increase cost and latency.", "Group programmatically, batch calls, skip unchanged segments, use cheaper model for extraction and stronger model only for repair."),
        ("Git awkwardness", "Raw data and drafts do not belong in Git.", "Use Git only for approved canonical pages; DB/object storage for everything else."),
        ("Sensitive data risk", "Users need sensitive data but not everyone should see it.", "Preserve data and filter retrieval by department, role, clearance, and sensitivity."),
        ("Local upload parsing", "PDFs, DOCX, and spreadsheets can be messy.", "Start with stable text parsers, skip OCR/scanned PDFs in v1, report unsupported files clearly."),
        ("Approval bottleneck", "Too many drafts may wait for review.", "Use confidence thresholds, bulk approval, and review only risky/conflicted pages later."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    set_table_width(table, [1.45, 2.2, 2.85])
    style_table(table, header_fill=BLUE_FILL)


def add_metrics(doc: Document) -> None:
    doc.add_heading("11. Metrics And Acceptance Criteria", level=1)
    table = doc.add_table(rows=1, cols=3)
    for idx, text in enumerate(["Metric", "V1 target", "Why it matters"]):
        table.rows[0].cells[idx].text = text
    rows = [
        ("fallback_page_count", "0", "No fake knowledge enters Cortex."),
        ("prompt_leakage_count", "0", "No prompt text becomes a knowledge page."),
        ("malformed_approved_pages", "0", "Approved pages are parseable and valid."),
        ("evidence_link_rate", ">= 95%", "Most claims point to source segments."),
        ("citation_correctness", ">= 85%", "Citations actually support the answer."),
        ("detail_retention_score", ">= 75%", "Cortex preserves details better than weak summaries."),
        ("hallucination_rate", "<= 5%", "Answers avoid unsupported claims."),
        ("time_to_first_searchable_result", "< 30 seconds", "Users can search before compilation completes."),
        ("unchanged_documents_skipped", ">= 80%", "Repeat syncs avoid wasted work."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    set_table_width(table, [2.05, 1.35, 3.1])
    style_table(table)


def add_test_plan(doc: Document) -> None:
    doc.add_heading("12. Test Plan", level=1)
    sections = [
        ("Schema tests", ["Creates source objects/documents/segments/relationships.", "Enforces tenant isolation.", "Stores connector metadata without schema failure."]),
        ("Connector tests", ["Local upload parses md, txt, pdf, docx, csv, xlsx, html.", "Notion recursively stores blocks as source segments.", "Google Docs preserves headings and tables.", "Slack groups threads and replies."]),
        ("Pipeline tests", ["Empty sync creates no fake pages.", "Fast raw segment indexing works before page compilation.", "Unchanged content hashes skip duplicate processing.", "Classifier/clusterer are not required for page creation."]),
        ("Validation tests", ["Rejects Input JSON, Expected Output, and </output_format> leakage.", "Rejects invalid YAML.", "Rejects propositions without evidence segment IDs.", "Rejected drafts are never committed to Git."]),
        ("Permission tests", ["Sales users see allowed sales data.", "Unauthorized users cannot retrieve restricted HR, finance, or executive propositions.", "Redacted evidence is not sent to the answer LLM."]),
        ("Evaluation tests", ["Compare Cortex versus baseline RAG on citation correctness, detail retention, hallucination rate, answer completeness, and speed."]),
    ]
    for heading, bullets in sections:
        doc.add_heading(heading, level=2)
        add_bullets(doc, bullets)


def add_future_platform(doc: Document) -> None:
    doc.add_heading("13. Future CLI, MCP, Sandbox, And Dynamic Tools", level=1)
    doc.add_paragraph(
        "The earlier architecture review recommended delaying powerful agent actions until Cortex can reliably ingest, cite, approve, and permission-filter knowledge. CLI and MCP should start as read-first interfaces over trusted knowledge. Dynamic tools and sandbox execution should come last."
    )

    doc.add_heading("CLI Roadmap", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Command"
    table.rows[0].cells[1].text = "Purpose"
    table.rows[0].cells[2].text = "When to ship"
    rows = [
        ("cortex init", "Initialize tenant/workspace config.", "After source model is stable."),
        ("cortex ingest", "Run local upload or connector sync.", "After Local Upload and Notion adapters."),
        ("cortex query", "Ask cited questions from approved pages and raw evidence.", "After permission-aware retrieval."),
        ("cortex eval", "Run Cortex vs RAG benchmark.", "Before claiming Cortex beats RAG."),
        ("cortex export", "Export approved pages and citation reports.", "After approval/Git truth works."),
        ("cortex mcp serve", "Expose MCP tools over the same backend API.", "After read-first MCP tools are safe."),
    ]
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    set_table_width(table, [1.55, 3.05, 1.9])
    style_table(table)

    doc.add_heading("MCP Tool Surface", level=2)
    add_bullets(
        doc,
        [
            "Start with read-first tools: query_knowledge, get_page, list_pages, get_source_segment, submit_feedback.",
            "Expose ingest tools only with explicit tenant permissions and job status tracking.",
            "Do not expose execute_tool until sandboxing, approval, and audit logs exist.",
            "Make MCP call the same backend APIs as the frontend and CLI; do not duplicate business logic.",
        ],
    )

    doc.add_heading("Sandbox Architecture", level=2)
    doc.add_paragraph(
        "Sandboxing should not be a thin subprocess wrapper. The MVP should use a separate execution service with per-run temporary workspace, timeout, CPU/memory limits, network policy, no inherited environment, explicit capability allowlist, ephemeral credentials, and full audit logs. Docker/rootless containers are the practical MVP. Restricted subprocess is development-only."
    )
    add_code(
        doc,
        "tool request -> policy check -> sandbox workspace -> capability allowlist\n"
        "             -> execute with timeout/resource/network limits\n"
        "             -> capture logs/artifacts -> approval/audit record",
    )

    doc.add_heading("Dynamic Skills And Tools", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Requirement"
    table.rows[0].cells[1].text = "Reason"
    rows = [
        ("Schema and manifest", "Generated skills/tools need declared inputs, outputs, owner, permissions, and version."),
        ("Source page link", "Every generated skill should trace back to approved Cortex knowledge."),
        ("Tests before publish", "A tool without tests should remain draft-only."),
        ("Sandbox trial run", "New tools must prove behavior in a restricted environment."),
        ("Human approval", "No newly generated tool should auto-run directly from retrieval."),
        ("Signed/pinned version", "Agents should execute an approved artifact, not mutable text."),
    ]
    for values in rows:
        cells = table.add_row().cells
        cells[0].text = values[0]
        cells[1].text = values[1]
    set_table_width(table, [2.2, 4.3])
    style_table(table, header_fill=BLUE_FILL)


def add_prompts(doc: Document) -> None:
    doc.add_heading("14. AI Coding Assistant Prompts", level=1)
    prompts = [
        (
            "Prompt 1: Normalized source schema",
            "Implement normalized source storage for Cortex. Add SQLite tables for source_objects, source_documents, source_segments, source_relationships, propositions, knowledge_page_drafts, and sync_runs. Preserve tenant isolation. Add dynamic migrations compatible with the current database initialization style. Include tests that create and query each entity. Do not change Git behavior yet.",
        ),
        (
            "Prompt 2: Local Upload adapter",
            "Implement a LocalUploadAdapter with discover, fetch, extract, and normalize methods. Support .md, .txt, text PDFs, .docx, .csv, .xlsx, and .html. Use pypdf or pdfplumber for PDFs, python-docx for DOCX, pandas/openpyxl for tables, and BeautifulSoup for HTML. Store parsed content as SourceDocument and SourceSegment records. Skip OCR/scanned PDFs in v1 with a clear unsupported reason.",
        ),
        (
            "Prompt 3: Notion migration",
            "Migrate Notion ingestion from raw message dictionaries to normalized source records. Keep discovery metadata, recursively fetch block children, preserve headings, lists, code blocks, tables, and child pages, and store every block or section as SourceSegment evidence. Empty sync must fail loudly and create no pages.",
        ),
        (
            "Prompt 4: Google Docs adapter",
            "Implement a GoogleDocsAdapter that normalizes document title, headings, paragraphs, tables, comments if available, and source URL into SourceObject, SourceDocument, SourceSegment, and SourceRelationship records. Use document structure as the segmentation strategy.",
        ),
        (
            "Prompt 5: Slack adapter",
            "Implement a SlackAdapter that fetches channels, threads, messages, replies, authors, timestamps, and file metadata. Normalize each thread as a SourceDocument and each message as a SourceSegment. Preserve channel, thread_ts, reactions, author, and timestamp in metadata_json.",
        ),
        (
            "Prompt 6: Replace page authorship path",
            "Replace the sentence-cluster page authoring path with a compiler that groups source segments by document structure or conversation thread. Classifier and clusterer must become optional enrichment steps, not required steps for page creation.",
        ),
        (
            "Prompt 7: Evidence-linked propositions",
            "Implement proposition extraction from grouped source segments. Every proposition must include evidence_segment_ids and source quote snippets. Prefer programmatic grouping before LLM calls. Reject propositions without evidence.",
        ),
        (
            "Prompt 8: Remove synthesis fallback pages",
            "Remove fallback Markdown page generation from the synthesizer. On LLM failure, store a rejected draft with error reason and source IDs. Never write or commit fallback pages to Git.",
        ),
        (
            "Prompt 9: Approval to Git",
            "Implement approval endpoints for drafts. Only approved drafts should be written as Markdown files and committed to the tenant Git repo. Rejected and draft outputs stay in SQLite/storage and are not queryable as truth by default.",
        ),
        (
            "Prompt 10: Permission-aware retrieval",
            "Update query retrieval to filter evidence by user department, role, clearance_level, source department, segment access_level, proposition sensitivity, and source permissions. Sensitive data is not removed at ingestion; it is excluded or redacted before answer generation if the user lacks clearance.",
        ),
        (
            "Prompt 11: Eval harness",
            "Build an evaluation harness that compares Cortex against baseline RAG on the same source documents and questions. Track citation correctness, detail retention, hallucination rate, answer completeness, permission failures, latency, and repeat-sync skip rate.",
        ),
    ]
    for title, prompt in prompts:
        doc.add_heading(title, level=2)
        add_code(doc, prompt)


def add_appendix(doc: Document) -> None:
    doc.add_heading("15. Appendix: Key Implementation Notes", level=1)
    doc.add_heading("Recommended Backend Shape", level=2)
    add_code(
        doc,
        "app/ingestion/connectors/base.py\n"
        "app/ingestion/connectors/local_upload.py\n"
        "app/ingestion/connectors/notion.py\n"
        "app/ingestion/connectors/google_docs.py\n"
        "app/ingestion/connectors/slack.py\n"
        "app/ingestion/source_store.py\n"
        "app/ingestion/compiler.py\n"
        "app/ingestion/propositions.py\n"
        "app/ingestion/drafts.py\n"
        "app/ingestion/validation.py\n"
        "app/retrieval/permissions.py\n"
        "app/retrieval/hybrid_query.py",
    )
    doc.add_heading("Definition Of Done", level=2)
    add_bullets(
        doc,
        [
            "Local Upload, Notion, Google Docs, and Slack produce normalized source records.",
            "Raw source segments become searchable before page compilation completes.",
            "Generated pages include valid YAML, sources, propositions, and evidence links.",
            "Invalid or prompt-leaked output becomes a rejected draft and is not committed.",
            "Approved pages are the only files committed to Git.",
            "Queries return answer, citations, pages read, source segments read, redactions, confidence, and knowledge gaps.",
            "Evaluation demonstrates stronger citation correctness and detail retention than baseline RAG.",
        ],
    )


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_document_styles(doc)
    add_title(doc)
    doc.add_page_break()
    add_contents(doc)
    doc.add_page_break()
    add_executive_summary(doc)
    add_previous_review_context(doc)
    add_qa_context(doc)
    add_changes(doc)
    add_how_it_works(doc)
    add_architecture(doc)
    add_data_model(doc)
    add_connector_plan(doc)
    add_implementation_roadmap(doc)
    add_downsides(doc)
    add_metrics(doc)
    add_test_plan(doc)
    add_future_platform(doc)
    add_prompts(doc)
    add_appendix(doc)
    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    build()
